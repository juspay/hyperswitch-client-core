#!/usr/bin/env python3
"""
Generate card-installments.docx from the markdown content.
This script uses python-docx to create a Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_code_block(doc, text, language="rescript"):
    """Add a code block with monospace formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('Card Installments — Implementation Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph()
    doc.add_paragraph('Branch: card-installments-mobile')
    doc.add_paragraph('Repo: hyperswitch-client-core')
    doc.add_paragraph('Reference PR (web SDK): hyperswitch-web#1412')
    doc.add_paragraph('Submodule PR: hyperswitch-sdk-utils#47')
    doc.add_paragraph()
    
    # Overview
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(
        'This feature allows customers to split a card payment into multiple installments. '
        'When a merchant\'s backend returns installment options for a payment intent, the UI presents:'
    )
    doc.add_paragraph('1. A checkbox — "Pay in installments"', style='List Number')
    doc.add_paragraph('2. A scrollable list of available plans (each showing the per-installment amount, interest rate, and total payable)', style='List Number')
    doc.add_paragraph('3. Validation — if the checkbox is checked but no plan is selected, an inline error is shown before confirm is allowed', style='List Number')
    doc.add_paragraph()
    doc.add_paragraph(
        'The checkbox is gated behind 6+ card digits for new card entry (matching the web SDK behaviour). '
        'For saved cards, it appears immediately when a card token is selected.'
    )
    
    # Architecture Decisions
    doc.add_heading('Architecture Decisions', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '#'
    hdr_cells[1].text = 'Decision'
    
    decisions = [
        ('AD-1', 'Reuse existing ClickableTextElement (checkbox) and CustomRadioButton (radio) components'),
        ('AD-2', 'Thread installment state via props — no new React Context'),
        ('AD-3', 'Flat locale strings with {placeholder} replacement via replaceLocaleParams'),
        ('AD-5', 'Gate installment UI on 6+ card digits for new card entry'),
        ('AD-6', 'installment_data?: JSON.t on redirectType — serialised at the JS layer'),
        ('AD-7', 'No native layer changes needed — the JS layer POSTs the confirm body directly'),
        ('AD-8', 'ScrollView with nestedScrollEnabled={true} for plan list'),
    ]
    
    for num, desc in decisions:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = desc
    
    doc.add_paragraph()
    
    # Files Changed
    doc.add_heading('Files Changed', level=1)
    
    # File 1
    doc.add_heading('1. src/types/AllApiDataTypes/AccountPaymentMethodType.res', level=2)
    doc.add_paragraph(
        'What changed: Added 4 new types and 4 parsers to represent installment data from the '
        'Payment Methods List (PML) API response. Added intent_data: option<intentData> to accountPaymentMethods.'
    )
    doc.add_heading('New Types', level=3)
    add_code_block(doc, '''type installmentAmountDetails = {
  amount_per_installment: float,
  total_amount: float,
}

type installmentPlan = {
  interest_rate: float,
  number_of_installments: int,
  billing_frequency: string,
  amount_details: installmentAmountDetails,
}

type installmentOption = {
  payment_method: string,
  available_plans: array<installmentPlan>,
}

type intentData = {
  installment_options: option<array<installmentOption>>,
  currency: string,
  amount: float,
}''')
    
    doc.add_heading('New Parsers', level=3)
    doc.add_paragraph('• parseAmountDetails — parses amount_per_installment and total_amount as floats')
    doc.add_paragraph('• parseInstallmentPlan — parses a single plan; number_of_installments decoded as float then Int.fromFloat')
    doc.add_paragraph('• parseInstallmentOption — parses payment_method + available_plans array')
    doc.add_paragraph('• parseIntentData — parses installment_options, currency, amount from the top-level PML response')
    
    # File 2
    doc.add_heading('2. src/types/AllApiDataTypes/PaymentConfirmTypes.res', level=2)
    doc.add_paragraph(
        'What changed: Added installment_data?: JSON.t to the redirectType record (the confirm request body).'
    )
    add_code_block(doc, '''type redirectType = {
  ...
  installment_data?: JSON.t,  // Added
}''')
    doc.add_paragraph(
        'This field carries the serialised installment selection to the Hyperswitch API. '
        'No native layer changes are needed — the JS layer makes the HTTP POST directly.'
    )
    
    # File 3
    doc.add_heading('3. src/utility/logics/Utils.res', level=2)
    doc.add_paragraph('What changed: Added 2 utility functions.')
    add_code_block(doc, '''let formatAmountWithTwoDecimals = (amount: float) =>
  amount->Float.toFixed(~digits=2)

let replaceLocaleParams = (template, params: array<(string, string)>) =>
  params->Array.reduce(template, (acc, (key, value)) =>
    acc->String.replaceAll("{" ++ key ++ "}", value)
  )''')
    doc.add_paragraph('• formatAmountWithTwoDecimals — formats currency amounts to 2 decimal places')
    doc.add_paragraph('• replaceLocaleParams — replaces {placeholder} tokens in locale strings')
    
    # File 4
    doc.add_heading('4. src/utility/logics/PaymentUtils.res', level=2)
    doc.add_paragraph('What changed: Added 2 functions and updated both confirm body generators.')
    doc.add_heading('New Functions', level=3)
    add_code_block(doc, '''let filterInstallmentPlansByPaymentMethod = (
  installmentOptions: array<installmentOption>,
  paymentMethod: string,
) =>
  installmentOptions
  ->Array.find(option => option.payment_method === paymentMethod)
  ->Option.map(option => option.available_plans)
  ->Option.getOr([])''')
    doc.add_paragraph(
        'Finds the installmentOption whose payment_method matches (e.g. "card") and returns its plans.'
    )
    add_code_block(doc, '''let installmentBody = (plan: option<installmentPlan>) =>
  plan->Option.map(p =>
    [
      ("number_of_installments", p.number_of_installments->Int.toFloat->JSON.Encode.float),
      ("billing_frequency", p.billing_frequency->JSON.Encode.string),
    ]
    ->Dict.fromArray
    ->JSON.Encode.object
  )''')
    doc.add_paragraph(
        'Serialises the selected plan to JSON for the confirm body. Matches the web SDK\'s installmentBody function.'
    )
    
    # File 5 - Locale
    doc.add_heading('5. shared-code/sdk-utils/types/LocaleDataType.res (submodule — PR #47)', level=2)
    doc.add_paragraph(
        'What changed: Added 7 locale string fields to the localeStrings type and default values.'
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Default Value'
    
    locales = [
        ('installmentPayInInstallments', 'Pay in installments'),
        ('installmentChoosePlan', 'Choose an installment plan'),
        ('installmentInterestFree', 'Interest free'),
        ('installmentInterestRate', '{rate}% interest'),
        ('installmentTotalPayable', 'Total'),
        ('installmentSelectPlanError', 'Please select an installment plan'),
        ('installmentPaymentLabel', '{count}x {amount}'),
    ]
    
    for field, default in locales:
        row_cells = table.add_row().cells
        row_cells[0].text = field
        row_cells[1].text = default
    
    # File 6
    doc.add_heading('6. src/hooks/S3ApiHook.res', level=2)
    doc.add_paragraph(
        'What changed: Added parsers for the 7 new locale fields inside getLocaleStrings. '
        'Each uses Utils.getString with the corresponding defaultLocale value as fallback.'
    )
    
    # File 7
    doc.add_heading('7. src/components/elements/InstallmentOptionItem.res (new file)', level=2)
    doc.add_paragraph('A single row in the installment plan list.')
    doc.add_heading('Props', level=3)
    doc.add_paragraph('• plan: installmentPlan')
    doc.add_paragraph('• currency: string')
    doc.add_paragraph('• isSelected: bool')
    doc.add_paragraph('• onSelect: unit => unit')
    doc.add_paragraph('• isLastItem: bool — suppresses the bottom divider on the last row')
    doc.add_heading('Layout', level=3)
    add_code_block(doc, '''[ Radio ] | {count}x {currency}{amount}    Total
           | {interest text}               {currency} {total}''', 'text')
    doc.add_paragraph(
        'Top row: payment label (bold) on the left, "Total" label (light) on the right. '
        'Bottom row: interest text (light) on the left, total amount (bold) on the right.'
    )
    
    # File 8
    doc.add_heading('8. src/components/elements/InstallmentOptions.res (new file)', level=2)
    doc.add_paragraph('Container component that wraps the checkbox and plan list.')
    doc.add_heading('Behaviour', level=3)
    doc.add_paragraph('• Hidden entirely if filterInstallmentPlansByPaymentMethod returns 0 plans')
    doc.add_paragraph('• Renders a ClickableTextElement checkbox labelled installmentPayInInstallments')
    doc.add_paragraph('• When checkbox is checked → shows plan list; when unchecked → clears selection and error')
    doc.add_paragraph('• Plan list is a ScrollView inside a themed bordered box')
    doc.add_paragraph('• Error text rendered below if errorString is non-empty')
    
    # File 9
    doc.add_heading('9. src/components/dynamic/TabElement.res', level=2)
    doc.add_paragraph(
        'What changed: Added optional ~onFormDataChange=_ => () prop. '
        'Called inside setFormData whenever form data updates.'
    )
    add_code_block(doc, '''let make = (
  ~isScreenFocus,
  ~processRequest,
  ~setConfirmButtonData,
  ~onFormDataChange=_ => (),   // Added
) => {
  ...
  let setFormData = React.useCallback1(data => {
    formDataRef->Option.map(ref => ref.current = data)->ignore
    setFormData(_ => data)
    onFormDataChange(data)     // Added
  }, [setFormData])''')
    doc.add_paragraph('This allows PaymentMethod.res to observe card number changes without adding a new context.')
    
    # File 10
    doc.add_heading('10. src/pages/payment/PaymentMethod.res', level=2)
    doc.add_paragraph('What changed: Full installment integration for new card payments (TAB flow).')
    doc.add_heading('State Added', level=3)
    add_code_block(doc, '''let (showInstallments, setShowInstallments) = React.useState(_ => false)
let (selectedInstallmentPlan, setSelectedInstallmentPlan) = React.useState(_ => None)
let (installmentsError, setInstallmentsError) = React.useState(_ => "")
let (cardDigitCount, setCardDigitCount) = React.useState(_ => 0)''')
    
    doc.add_heading('Card Digit Tracking', level=3)
    doc.add_paragraph(
        'onFormDataChange callback traverses the nested React Final Form values dict to extract the card number:'
    )
    add_code_block(doc, '''let onFormDataChange = (data: Dict.t<JSON.t>) => {
  let cardNumber =
    data
    ->Dict.get("payment_method_data")
    ->Option.flatMap(JSON.Decode.object)
    ->Option.flatMap(d => d->Dict.get("card"))
    ->Option.flatMap(JSON.Decode.object)
    ->Option.flatMap(d => d->Dict.get("card_number"))
    ->Option.flatMap(JSON.Decode.string)
    ->Option.getOr("")
    ->Validation.clearSpaces
  setCardDigitCount(_ => cardNumber->String.length)
}''')
    doc.add_paragraph(
        'Note: React Final Form stores dotted field names as nested objects. '
        'payment_method_data.card.card_number becomes a nested dict, not a flat key.'
    )
    
    # File 11
    doc.add_heading('11. src/components/dynamic/DynamicComponent.res', level=2)
    doc.add_paragraph(
        'What changed: Same installment integration as PaymentMethod.res, adapted for the widget/button flow.'
    )
    doc.add_paragraph(
        'Key Difference: cardDigitCount is derived via React.useMemo1 directly from local formData state.'
    )
    
    # File 12
    doc.add_heading('12. src/pages/payment/SavedPaymentSheet.res', level=2)
    doc.add_paragraph(
        'What changed: Installment integration for saved card tokens. '
        'No digit gate — the installment UI appears as soon as a card token is selected.'
    )
    doc.add_paragraph('Validation Gate: Installment check runs before the existing CVV validation.')
    
    # Data Flow
    doc.add_page_break()
    doc.add_heading('Data Flow', level=1)
    doc.add_paragraph()
    add_code_block(doc, '''PML API response (/account/payment_methods)
  └─ intent_data.installment_options[]
       └─ AccountPaymentMethodType.parseIntentData
            └─ accountPaymentMethods.intent_data
                 └─ PaymentMethod / DynamicComponent / SavedPaymentSheet
                      └─ installmentOptions (array<installmentOption>)
                           └─ InstallmentOptions component
                                └─ filterInstallmentPlansByPaymentMethod
                                     └─ InstallmentOptionItem (per plan)

User selects plan → selectedInstallmentPlan: option<installmentPlan>

Confirm button pressed
  └─ installmentBody(selectedInstallmentPlan) → JSON.t
       └─ generateCardConfirmBody / generateSavedCardConfirmBody
            └─ redirectType.installment_data
                 └─ HTTP POST /payments/{id}/confirm''', 'text')
    
    # API Contract
    doc.add_heading('API Contract', level=1)
    doc.add_heading('Input (from PML response)', level=2)
    add_code_block(doc, '''{
  "intent_data": {
    "currency": "BRL",
    "amount": 10000.0,
    "installment_options": [
      {
        "payment_method": "card",
        "available_plans": [
          {
            "interest_rate": 0.0,
            "number_of_installments": 3,
            "billing_frequency": "month",
            "amount_details": {
              "amount_per_installment": 3333.33,
              "total_amount": 10000.0
            }
          }
        ]
      }
    ]
  }
}''', 'json')
    
    doc.add_heading('Output (in confirm body)', level=2)
    add_code_block(doc, '''{
  "installment_data": {
    "number_of_installments": 3,
    "billing_frequency": "month"
  }
}''', 'json')
    
    # UI Behaviour
    doc.add_heading('UI Behaviour Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Condition'
    hdr_cells[1].text = 'Installment UI'
    
    behaviors = [
        ('No installment_options in PML response', 'Hidden (0 plans)'),
        ('New card, < 6 digits entered', 'Hidden'),
        ('New card, ≥ 6 digits entered', 'Checkbox visible'),
        ('Saved card token selected (card type)', 'Checkbox visible'),
        ('Checkbox unchecked', 'Plan list hidden'),
        ('Checkbox checked', 'Plan list + "Choose a plan" label visible'),
        ('Plan selected', 'Radio filled, no error'),
        ('Confirm pressed, checkbox checked, no plan selected', 'Inline error shown'),
    ]
    
    for condition, result in behaviors:
        row_cells = table.add_row().cells
        row_cells[0].text = condition
        row_cells[1].text = result
    
    # Build & Deploy
    doc.add_heading('Build & Deploy', level=1)
    doc.add_paragraph('Changes to .res files require all three steps:')
    add_code_block(doc, '''# 1. Compile ReScript
npx rescript build

# 2. Bundle for Android
npx react-native bundle \\
  --reset-cache \\
  --platform android \\
  --dev false \\
  --entry-file index.js \\
  --bundle-output android/app/src/main/assets/hyperswitch.bundle

# 3. Install APK
yarn android''', 'bash')
    doc.add_paragraph('The Android app loads from the pre-bundled asset file — Metro live reload does not apply.')
    
    # Related PRs
    doc.add_heading('Related PRs', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Repo'
    hdr_cells[1].text = 'PR'
    hdr_cells[2].text = 'Description'
    
    prs = [
        ('hyperswitch-sdk-utils', '#47', 'Adds locale types for installments to LocaleDataType.res'),
        ('hyperswitch-web', '#1412', 'Reference implementation (web SDK)'),
    ]
    
    for repo, pr, desc in prs:
        row_cells = table.add_row().cells
        row_cells[0].text = repo
        row_cells[1].text = pr
        row_cells[2].text = desc
    
    # Deferred
    doc.add_heading('Deferred / Out of Scope', level=1)
    doc.add_paragraph('• Payment method config (paymentMethodsConfig, showInterestRates, Terms component) — tracked separately')
    doc.add_paragraph('• Native layer changes (iOS/Android) — not required; JS layer POSTs the confirm body directly')
    
    # Save document
    doc.save('card-installments.docx')
    print("Generated: card-installments.docx")

if __name__ == '__main__':
    main()
